import os
import io
import shutil
import tempfile
import logging
from typing import List, Optional
from contextlib import asynccontextmanager
from fastapi import FastAPI, UploadFile, File, HTTPException, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("lexagent.backend.main")

from parsers.document_parser import parse_pdf, parse_docx, split_clauses, compare_documents, generate_text_diff
from agents.risk_analyzer import analyze_clause_risk
from agents.negotiator import run_negotiation_cycle

@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup logic
    logger.info("Initializing LexAgent Backend...")
    try:
        from rag.ingestion import seed_database
        seed_database()
        logger.info("ChromaDB seeded successfully.")
    except Exception as e:
        logger.error(f"Failed to seed ChromaDB on startup: {e}")
    yield
    # Shutdown logic (no-op)
    logger.info("Shutting down LexAgent Backend...")

app = FastAPI(
    title="LexAgent API",
    description="Backend API for AI-powered legal document risk analysis and negotiation",
    version="1.0.0",
    lifespan=lifespan
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:5174", "http://localhost:5175", "http://127.0.0.1:5173", "http://127.0.0.1:5174"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
# CORS Middleware Setup


# Request schemas
class ClauseInput(BaseModel):
    id: str
    title: str
    original_text: str
    counterparty_text: Optional[str] = ""
    has_diff: Optional[bool] = False
    diff: Optional[str] = ""

class AnalyzeRequest(BaseModel):
    clauses: List[ClauseInput]

class NegotiateRequest(BaseModel):
    clause_id: str
    clause_text: str  # The text currently being negotiated
    original_text: Optional[str] = ""
    iteration: int
    user_feedback: str  # "reject" | "accept" | "try_again"

class ExportClause(BaseModel):
    title: str
    text: str

class ExportRequest(BaseModel):
    clauses: List[ExportClause]

@app.get("/api/health")
async def health_check():
    """Health check endpoint to ensure API service is operational."""
    return {
        "status": "healthy",
        "openai_key_configured": bool(os.environ.get("OPENAI_API_KEY")),
        "chroma_dir": os.environ.get("CHROMA_PERSIST_DIR", "./chroma_db")
    }

@app.post("/api/upload")
async def upload_documents(
    files: List[UploadFile] = File(...)
):
    """
    Upload 1 or 2 contract files (PDF or DOCX).
    If 1 file is uploaded, returns parsed clauses.
    If 2 files are uploaded, diffs them and returns compared clauses.
    """
    if not files or len(files) > 2:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="You must upload exactly 1 or 2 documents."
        )
        
    temp_files = []
    try:
        # Save files to temp directory to allow fitz/python-docx to read them
        for f in files:
            suffix = os.path.splitext(f.filename)[1].lower()
            if suffix not in [".pdf", ".docx"]:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Unsupported file format '{suffix}'. Only PDF and DOCX are allowed."
                )
            
            temp_f = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
            shutil.copyfileobj(f.file, temp_f)
            temp_f.close()
            temp_files.append((temp_f.name, suffix))

        # Parse text from files
        parsed_texts = []
        for path, ext in temp_files:
            if ext == ".pdf":
                parsed_texts.append(parse_pdf(path))
            else:
                parsed_texts.append(parse_docx(path))

        # Generate clauses
        if len(parsed_texts) == 1:
            clauses = split_clauses(parsed_texts[0])
            # Format return clauses to match full compared schema structure
            formatted_clauses = []
            for c in clauses:
                formatted_clauses.append({
                    "id": c["id"],
                    "title": c["title"],
                    "original_text": c["text"],
                    "counterparty_text": "",
                    "has_diff": False,
                    "diff": ""
                })
            return {"mode": "single", "clauses": formatted_clauses}
        else:
            clauses1 = split_clauses(parsed_texts[0])
            clauses2 = split_clauses(parsed_texts[1])
            compared = compare_documents(clauses1, clauses2)
            return {"mode": "compare", "clauses": compared}

    except Exception as e:
        logger.error(f"Error during document upload/parse: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Document parsing failed: {str(e)}"
        )
    finally:
        # Clean up temp files
        for path, _ in temp_files:
            if os.path.exists(path):
                os.remove(path)

@app.post("/api/analyze")
async def analyze_clauses(request: AnalyzeRequest):
    """
    Runs risk analysis on all provided clauses.
    Returns the clauses with risk level, explanation, favors, and RAG context.
    """
    logger.info(f"Analyzing risks for {len(request.clauses)} clauses")
    analyzed_clauses = []
    
    for c in request.clauses:
        # Use counterparty text if it exists (meaning they edited it), else use original
        text_to_analyze = c.counterparty_text if (c.counterparty_text and c.has_diff) else c.original_text
        if not text_to_analyze.strip():
            # Fallback if both empty
            text_to_analyze = c.original_text
            
        try:
            analysis = analyze_clause_risk(text_to_analyze)
            analyzed_clauses.append({
                "id": c.id,
                "title": c.title,
                "original_text": c.original_text,
                "counterparty_text": c.counterparty_text,
                "has_diff": c.has_diff,
                "diff": c.diff,
                "risk_level": analysis.get("risk_level", "LOW"),
                "explanation": analysis.get("explanation", ""),
                "favors": analysis.get("favors", "neutral"),
                "rag_context": analysis.get("rag_context", [])
            })
        except Exception as e:
            logger.error(f"Error analyzing clause {c.id}: {e}")
            analyzed_clauses.append({
                **c.model_dump(),
                "risk_level": "MEDIUM",
                "explanation": f"Error during analysis: {str(e)}",
                "favors": "neutral",
                "rag_context": []
            })
            
    return {"clauses": analyzed_clauses}

@app.post("/api/negotiate")
async def negotiate_clause(request: NegotiateRequest):
    """
    Starts or continues the negotiation agent loop for a clause.
    Returns the counter-proposal, diff, rationale, iteration, and status.
    """
    logger.info(f"Negotiate called on clause {request.clause_id}, iteration {request.iteration}, feedback: '{request.user_feedback}'")
    
    # Run the cycle
    original = request.original_text or request.clause_text
    
    try:
        final_state = run_negotiation_cycle(
            clause_id=request.clause_id,
            original_text=original,
            current_text=request.clause_text,
            iteration=request.iteration,
            user_feedback=request.user_feedback
        )
        
        proposed = final_state.get("proposed_text", "")
        # Compute diff between current text and proposed text
        diff_val = generate_text_diff(request.clause_text, proposed)
        
        return {
            "proposed_text": proposed,
            "diff": diff_val,
            "rationale": final_state.get("rationale", ""),
            "iteration": final_state.get("iteration", request.iteration + 1),
            "status": final_state.get("status", "negotiating")
        }
    except Exception as e:
        logger.error(f"Negotiation error: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Negotiation cycle failed: {str(e)}"
        )

@app.post("/api/export")
async def export_contract(request: ExportRequest):
    """
    Assembles final negotiated/original clauses into a clean Word document.
    Returns it as a downloadable attachment.
    """
    logger.info(f"Exporting contract with {len(request.clauses)} clauses")
    
    try:
        import docx
        doc = docx.Document()
        
        # Add basic document formatting
        title = doc.add_heading("LEXAGENT NEGOTIATED CONTRACT", level=0)
        title.alignment = 1  # Center alignment
        
        doc.add_paragraph("This document compiles the finalized and negotiated contract terms processed by LexAgent AI. Below are the agreed-upon clauses.")
        doc.add_paragraph().paragraph_format.space_after = docx.shared.Pt(24)
        
        for c in request.clauses:
            # Add clause title as Heading 2
            h = doc.add_heading(c.title, level=2)
            h.paragraph_format.space_before = docx.shared.Pt(18)
            h.paragraph_format.space_after = docx.shared.Pt(6)
            
            # Add clause body text
            p = doc.add_paragraph(c.text)
            p.paragraph_format.space_after = docx.shared.Pt(12)
            p.paragraph_format.line_spacing = 1.15
            
        # Write to byte stream
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        headers = {
            'Content-Disposition': 'attachment; filename="LexAgent_Final_Contract.docx"'
        }
        
        return StreamingResponse(
            file_stream,
            headers=headers,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        logger.error(f"Export error: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Exporting document failed: {str(e)}"
        )

if __name__ == "__main__":
    import uvicorn
    # When running main.py directly, start the uvicorn server
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
