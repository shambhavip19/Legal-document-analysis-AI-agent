import docx
import os

def create_original_contract():
    doc = docx.Document()
    doc.add_heading("SOFTWARE DEVELOPMENT AND SERVICES AGREEMENT", level=1)
    doc.add_paragraph("This Software Development and Services Agreement (the 'Agreement') is entered into by and between Acme Corp ('Customer') and DevForce LLC ('Provider').")
    
    # 1. Preamble
    doc.add_heading("Section 1: Scope of Services", level=2)
    doc.add_paragraph("Provider shall perform the software engineering, template construction, and implementation services specified in Statement of Work (SOW) 1 attached hereto. All work shall be performed in a professional manner.")

    # 2. Indemnification (High Risk - Unilateral)
    doc.add_heading("Section 2: Indemnification", level=2)
    doc.add_paragraph("Provider agrees to defend, indemnify, and hold harmless Customer and its employees, directors, and agents from and against any and all claims, losses, damages, liabilities, and expenses (including attorneys' fees) arising out of or related to the Services, any deliverables provided hereunder, or any breach of this Agreement by Provider. Customer shall have no reciprocal obligation to indemnify Provider.")

    # 3. Limitation of Liability (High Risk - No Cap / Exclusions)
    doc.add_heading("Section 3: Limitation of Liability", level=2)
    doc.add_paragraph("UNDER NO CIRCUMSTANCES SHALL CUSTOMER BE LIABLE TO PROVIDER FOR ANY INDIRECT, SPECIAL, OR CONSEQUENTIAL DAMAGES ARISING OUT OF THIS AGREEMENT. PROVIDER'S LIABILITY UNDER THIS AGREEMENT IS UNLIMITED AND SHALL NOT BE CAPPED BY THE VALUE OF THE FEES PAID UNDER THIS AGREEMENT.")

    # 4. Intellectual Property (High Risk - Transfer of pre-existing)
    doc.add_heading("Section 4: Intellectual Property Ownership", level=2)
    doc.add_paragraph("All deliverables, software, code, designs, documentation, and work product developed under this Agreement shall be considered 'works made for hire' and shall belong exclusively to Customer upon creation. Provider hereby assigns all right, title, and interest in and to the deliverables, including all patent, copyright, and trade secret rights therein, to Customer. Provider further agrees that any pre-existing code, templates, or background tools used in the deliverables shall also become the property of Customer.")

    # 5. Payment Terms (Net 90)
    doc.add_heading("Section 5: Payment and Invoice Terms", level=2)
    doc.add_paragraph("Customer shall pay all invoices submitted by Provider within ninety (90) days from receipt of the invoice. Late payments shall not bear interest and shall not constitute a material breach of this Agreement.")

    # 6. Governing Law (Unfamiliar Forum)
    doc.add_heading("Section 6: Governing Law and Venue", level=2)
    doc.add_paragraph("This Agreement shall be governed by, and construed in accordance with, the laws of the State of Alaska, without regard to its conflict of law principles. Any legal action arising under this Agreement shall be brought exclusively in the state or federal courts located in Anchorage, Alaska.")

    doc.save("sample_contract_original.docx")
    print("Created sample_contract_original.docx")

def create_counterparty_contract():
    doc = docx.Document()
    doc.add_heading("SOFTWARE DEVELOPMENT AND SERVICES AGREEMENT", level=1)
    doc.add_paragraph("This Software Development and Services Agreement (the 'Agreement') is entered into by and between Acme Corp ('Customer') and DevForce LLC ('Provider').")
    
    # 1. Preamble
    doc.add_heading("Section 1: Scope of Services", level=2)
    doc.add_paragraph("Provider shall perform the software engineering, template construction, and implementation services specified in Statement of Work (SOW) 1 attached hereto. All work shall be performed in a professional manner.")

    # 2. Indemnification (High Risk - Unilateral, but slightly reworded by counterparty)
    doc.add_heading("Section 2: Indemnification", level=2)
    doc.add_paragraph("Provider agrees to defend, indemnify, and hold harmless Customer and its employees, directors, and agents from and against any and all claims, losses, damages, liabilities, and expenses (including attorneys' fees) arising out of or related to the Services, any deliverables provided hereunder, or any breach of this Agreement by Provider. Customer shall have no reciprocal obligation to indemnify Provider. Provider's total indemnity liability shall be capped at $50,000.")

    # 3. Limitation of Liability (Counterparty modified to add a cap)
    doc.add_heading("Section 3: Limitation of Liability", level=2)
    doc.add_paragraph("UNDER NO CIRCUMSTANCES SHALL CUSTOMER BE LIABLE TO PROVIDER FOR ANY INDIRECT, SPECIAL, OR CONSEQUENTIAL DAMAGES. PROVIDER'S LIABILITY UNDER THIS AGREEMENT SHALL BE CAP-LIMITED TO THE AMOUNT OF FEWER THAN THE PAST SIX (6) MONTHS FEES OR $10,000.")

    # 4. Intellectual Property (Counterparty reworded)
    doc.add_heading("Section 4: Intellectual Property Ownership", level=2)
    doc.add_paragraph("All deliverables developed under this Agreement shall belong exclusively to Customer upon creation. Provider hereby assigns all right, title, and interest in and to the deliverables, including all patent, copyright, and trade secret rights therein, to Customer. Any pre-existing code, templates, or background tools used in the deliverables shall remain the property of Provider, who hereby grants Customer a non-exclusive license to use them.")

    # 5. Payment Terms (Counterparty changed Net 90 to Net 60)
    doc.add_heading("Section 5: Payment and Invoice Terms", level=2)
    doc.add_paragraph("Customer shall pay all invoices submitted by Provider within sixty (60) days from receipt of the invoice. Late payments shall bear interest at 0.5% per month.")

    # 6. Governing Law (Counterparty changed venue to England and Wales)
    doc.add_heading("Section 6: Governing Law and Venue", level=2)
    doc.add_paragraph("This Agreement shall be governed by, and construed in accordance with, the laws of England and Wales. Any legal action arising under this Agreement shall be brought exclusively in London, United Kingdom.")

    doc.save("sample_contract_counterparty.docx")
    print("Created sample_contract_counterparty.docx")

if __name__ == "__main__":
    create_original_contract()
    create_counterparty_contract()
